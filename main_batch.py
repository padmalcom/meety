import vosk
import wave
import json
from pydub import AudioSegment
from datetime import datetime
import os
import numpy as np
import pandas as pd
from spacy.lang.de import German
from tqdm import tqdm
import math
from spacy.lang.de import German 
from difflib import SequenceMatcher
import copy
import sys
from loguru import logger
from docx import Document
from docx.enum.text import WD_UNDERLINE
from timeit import default_timer as timer
from datetime import timedelta

import streamlit as st

from recasepunc import CasePuncPredictor

# todo set tags in text block

# This can be a word or a sentence
class AnnotatedSequence:
	def __init__(self, sequence, start, end, speaker, confidence, min_confidence):
		self.sequence = sequence
		self.start = start
		self.end = end
		self.speaker = speaker
		self.confidence = confidence
		self.min_confidence = min_confidence
		
	def get_formated_start(self):
		return self.timeString(self.start)
		
	def get_formated_end(self):
		return self.timeString(self.end)
		
	def timeString(self, seconds):
		minutes = seconds / 60
		seconds = seconds % 60
		hours = int(minutes / 60)
		minutes = int(minutes % 60)
		return '%i:%02i:%06.3f' % (hours, minutes, seconds)
		
	def __repr__(self):
		return "(sequence: {}, start: {}, end: {}, speaker: {}, confidence: {}, min_confidence: {}".format(self.sequence, self.get_formated_start(), self.get_formated_end(), self.speaker, self.confidence, self.min_confidence)

class Transcriber:
	
	def __init__(self):
		vosk.SetLogLevel(-2)
		model_path = 'vosk-model-small-de-0.15'
		#model_path = 'vosk-model-de-0.21'
		spk_model_path = 'vosk-model-spk-0.4'
		self.COSINE_DIST = 0.8
		self.UNDETECTED_SPEAKER = "undetected speaker"
		self.UNKNOWN_SPEAKER = "unknown speaker"
		self.sample_rate = 16000
		self.model = vosk.Model(model_path)
		self.spk_model = vosk.SpkModel(spk_model_path)
		self.rec = vosk.KaldiRecognizer(self.model, self.sample_rate, self.spk_model)
		
		self.speakers={}
		#self.speakers["jonas"] = [-0.849626, 0.1394, 0.377193, 0.04026, -0.516814, 0.684988, 0.276028, 0.44519, -0.071336, 0.236369, 1.438427, -1.16203, 0.171904, -0.246187, 1.162354, 1.754254, -0.870339, 0.581138, 0.277063, -0.213094, -0.178992, 1.307812, 0.383807, -1.808675, -0.183495, 0.595374, 0.552995, 1.130693, 1.509223, 0.55946, 0.377448, -0.874153, -0.805022, -0.053774, -0.024277, 0.100899, -0.16061, 0.979247, -0.806801, 1.075389, 0.054593, 0.457873, -0.123776, 0.102562, -1.03972, -0.405821, -2.296902, 1.552534, 1.987395, -0.178972, -1.699092, -1.007383, -0.044273, 0.668422, -0.387606, 1.032058, 1.218006, -0.004968, -2.110302, -0.588267, -1.630525, -0.486805, -0.149978, -0.323995, 0.057739, -1.591965, -0.507674, 2.15651, 0.087731, 0.380945, 1.659711, 1.078468, -0.96696, 1.021768, -0.429574, 1.221862, -1.227425, -0.760518, 0.883041, 1.018671, 1.070466, -0.274897, 0.163634, -0.747442, -0.05772, 0.477834, -2.581385, 0.522382, -1.506052, 1.325543, -0.888436, -0.042877, -0.999906, -0.163342, -0.045787, -2.195781, 1.082597, -0.170146, -0.511741, 0.051157, -1.417514, -1.123549, 0.37979, -0.715364, 0.868087, 0.182235, 1.849722, 2.322327, 0.353243, 0.963822, -1.506676, 1.008916, 0.883645, -0.744136, -0.141038, -0.231468, 0.424568, 0.640757, 1.176899, 0.703505, -1.75687, -0.789799, -0.169159, -1.980155, -0.769657, 0.923687, -0.850422, 1.737328]
		self.new_speaker_index = 0
		
		punc_predict_path = os.path.abspath('vosk-recasepunc-de-0.21\\checkpoint')
		self.casePuncPredictor = CasePuncPredictor(punc_predict_path, lang="de")
		logger.info("Initialization done.")
		
	def repair_text(self, text):
		tokens = list(enumerate(self.casePuncPredictor.tokenize(text)))
		text = ''
		for token, case_label, punc_label in self.casePuncPredictor.predict(tokens, lambda x: x[1]):
			mapped = self.casePuncPredictor.map_punc_label(self.casePuncPredictor.map_case_label(token[1], case_label), punc_label)
			logger.trace("Token {}, case_label {}, punc_label {}, mapped {}", token, case_label, punc_label, mapped)
			if token[1].startswith('##'):
				text += mapped
			else:
				text +=  ' ' + mapped
		return text.strip()
		
	def get_speaker(self, fingerprint):
		best_speaker = ""
		if len(fingerprint) == 0:
			return best_speaker
		min_cd = 1000
		for s in self.speakers.keys():
			cd = self.cosine_dist(self.speakers[s], fingerprint)
			if cd < min_cd and cd < self.COSINE_DIST:
					best_speaker = s
					min_cd = cd
		if best_speaker.strip() == "":
			self.speakers[self.UNKNOWN_SPEAKER+str(self.new_speaker_index)] = fingerprint
			best_speaker = self.UNKNOWN_SPEAKER+str(self.new_speaker_index)
			self.new_speaker_index += 1
		return best_speaker, min_cd

	def cosine_dist(self, x, y):
		nx = np.array(x)
		ny = np.array(y)
		return 1 - np.dot(nx, ny) / np.linalg.norm(nx) / np.linalg.norm(ny)
		
	def get_speaker_from_audio_segment(self, audio_file, start, stop):
		logger.debug("Split audio {} from {} to {}", audio_file, start, stop)
		
		self.rec = vosk.KaldiRecognizer(self.model, self.sample_rate, self.spk_model)
		self.rec.SetWords(False)
		audio = AudioSegment.from_wav(audio_file)
		excerpt = audio[int(start*1000):int(stop*1000)]
		logger.debug("Length of excerpt is {} ms.", len(excerpt))

		timestamp = datetime.now().microsecond
		excerpt.export(str(timestamp) + ".wav", format='wav')
		wf = wave.open(str(timestamp) + ".wav", "rb")
		while True:
		#for i in tqdm(range(0, math.ceil(wf.getnframes() / 4000))):
			data = wf.readframes(4000)
			if len(data) == 0:
				break
			if self.rec.AcceptWaveform(data):
				pass
		res = self.rec.FinalResult()
		jres = json.loads(res)
		wf.close()
		os.remove(str(timestamp) + ".wav")
		if 'spk' in jres:
			speaker_name, cd = self.get_speaker(jres['spk'])
			return speaker_name, cd
		else:
			return self.UNDETECTED_SPEAKER, 1000
	
	def get_words_from_text(self, audio_file):
		rec_results = []
		self.rec = vosk.KaldiRecognizer(self.model, self.sample_rate, self.spk_model)
		self.rec.SetWords(True)
		wf = wave.open(audio_file, "rb")
		for i in tqdm(range(0, math.ceil(wf.getnframes() / 4000))):
		#while True:
		   data = wf.readframes(4000)
		   if len(data) == 0:
			   break
		   if self.rec.AcceptWaveform(data):
			   rec_results.append(self.rec.Result())
		rec_results.append(self.rec.FinalResult())
		wf.close()
		word_results = []
		full_text = ""
		for i, res in enumerate(rec_results):
			words = json.loads(res).get('result')
			if not words:
				continue
			for w in words:
				word_results.append(AnnotatedSequence(w['word'], w['start'], w['end'], None, w['conf'], w['conf']))
				logger.trace("Found new word {}.", word_results[len(word_results)-1])
				full_text += ' ' + w['word']
				
		return word_results, full_text.strip()
				
	def best_sentence_fit(self, unrepaired_words, unrepaired_start_index, repaired_words): # new
		
		punctuation = '''!()-[]{};:'"\,<>./?@#$%^&*_~'''
		
		repaired_words_without_punct = [w for w in repaired_words if not w in punctuation]

		repaired_sentence_without_punct = ' '.join(repaired_words_without_punct)

		# create i sentences
		best_ratio = 0
		best_end_index = 0
		for i in range(unrepaired_start_index, len(unrepaired_words)):
			sub_list = [x.sequence for x in unrepaired_words[unrepaired_start_index:i+1]]
			sentence_to_compare = ' '.join(sub_list)
			ratio = SequenceMatcher(None, repaired_sentence_without_punct, sentence_to_compare).ratio()
			logger.trace("{}: sentence to compare: {} ratio: {}", i, sentence_to_compare, ratio)
			if ratio > best_ratio:
				best_ratio = ratio
				best_end_index = i
				
			# Experimental - don't compare sentences, if the length of the constructed length is twice the size
			if len(repaired_sentence_without_punct) < len(sentence_to_compare) * 2:
				break
		logger.debug("Best end {}, best ratio {}.", best_end_index, best_ratio)
		return best_end_index		
		
	def get_sentences(self, words, repaired_text):
		sentences = []
		
		# split repaired_text in sentences
		nlp = German()
		nlp.add_pipe('sentencizer')
		doc = nlp(repaired_text)
		repaired_text_sentences = [str(sent).strip() for sent in doc.sents]
		
		word_index = 0
		for rts in repaired_text_sentences:	
			words_in_sentence = [token.text for token in nlp(rts)]
			end_index = self.best_sentence_fit(words, word_index, words_in_sentence)
			confidences = [x.confidence for x in words[word_index:end_index]]
			avg_conf = 0
			min_conf = 0
			if len(confidences):
				avg_conf = sum(confidences) / (len(confidences) + 0.00000001)
				min_conf = min(confidences)
			
			#print("Lenth {}, start_index {}, end_index {}".format(len(words), word_index, end_index))
			
			sentences.append(AnnotatedSequence(rts, words[word_index].start, words[end_index].end, None, avg_conf, min_conf))
			word_index = end_index + 1
		return sentences
		
	# merge sentences with same speakers
	def merge_sentences(self, annotated_sentences):
		results = []
		current_speaker = ""
		current_sentence = ""
		current_start = 0
		current_end = 0
		confidences = []
		for i, anse in enumerate(annotated_sentences):
			if i == 0:
				current_speaker = anse.speaker
				current_sentence = anse.sequence
				current_start = anse.start
				current_end = anse.end
				confidences.append(anse.confidence)
				continue
		
			if anse.speaker != current_speaker and len(current_sentence) > 0:
				avg_conf = 0
				min_conf = 0			
				if len(confidences):
					avg_conf = sum(confidences) / (len(confidences) + 0.00000001)
					min_conf = min(confidences)
				results.append(AnnotatedSequence(current_sentence, current_start, current_end, current_speaker, avg_conf, min_conf))
				current_speaker = anse.speaker
				current_sentence = anse.sequence
				current_start = anse.start
				current_end = 0
				confidences.clear()
				confidences.append(anse.confidence)
			else:
				current_sentence += ' ' + anse.sequence
				current_end = anse.end
				confidences.append(anse.confidence)
		
		# Add last item if not empty
		if len(current_sentence) > 0:
			avg_conf = 0
			min_conf = 0		
			if len(confidences):
				avg_conf = sum(confidences) / (len(confidences) + 0.00000001)
				min_conf = min(confidences)		
			results.append(AnnotatedSequence(current_sentence, current_start, current_end, current_speaker, avg_conf, min_conf))
				
		return results
		
		
	def get_speakers_from_sentences(self, annotated_sentences, audio_file):
		results = []
		for s in tqdm(annotated_sentences):
			logger.debug("Getting speaker for {}.", s)
			speaker_name, distance = self.get_speaker_from_audio_segment(audio_file, s.start, s.end)
			logger.debug("Found speaker {}.", speaker_name)
			new_sentence = copy.copy(s)
			new_sentence.speaker = speaker_name
			results.append(new_sentence)
		return results
		
	def fix_audio(self, audio_file):
		file_name, file_extension = os.path.splitext(audio_file)
		audio = None
		if file_extension.lower() == '.mp3':
			audio = AudioSegment.from_mp3(audio_file)
		elif file_extension.lower() == '.wav':
			audio = AudioSegment.from_wav(audio_file)
		else:
			logger.error("File format {} not supported.", file_extension)
			sys.exit(0)
			
		audio = audio.set_channels(1)
		audio = audio.set_frame_rate(16000)
		audio = audio.set_sample_width(2)
		timestamp = datetime.now().microsecond
		file_name = "conv" + str(timestamp) + ".wav"
		audio.export(file_name, format='wav', bitrate="64k")
		return file_name
		
	def write_docx(self, annotated_sentences, out_file):
		document = Document()
		document.add_heading('Meeting', 0)
		
		#p = None
		#last_speaker = None
		for anse in annotated_sentences:
		#	if anse.speaker != last_speaker or p == None:
		#		p = document.add_paragraph('')
		#		p.add_run(anse.speaker + ":").bold = True

		#	if anse.min_confidence < 0.5:
		#		p.add_run(' ' + anse.sequence).underline = WD_UNDERLINE.WAVY
		#	else:
		#		p.add_run(' ' + anse.sequence).underline = False
		#	last_speaker = anse.speaker
			p = document.add_paragraph('')
			p.add_run(anse.speaker + ":").bold = True
			p.add_run(' ' + anse.sequence)
		document.save(out_file)
			
if __name__ == '__main__':
	logger.remove()
	logger.add(sys.stderr, level="INFO")
	
	st.title("Meety - Herzlichen Willkommen!")
	st.header("AI-basierte Audiotranskription f??r Web-Konferenzen")
	
	st.markdown("Diese Anwendung hilft dir, Meetings automatisiert zu verschriftlichen und dabei zwischen verschiedenen Sprechnern zu unterscheiden. Neugierig geworden? Probier es mal aus!")
	
	uploaded_file = st.file_uploader("W??hlen eine Audiodatei", type=['wav', 'mp3'])
	if uploaded_file is not None:
		bytes_data = uploaded_file.getvalue()

		start = timer()
		t = Transcriber()
		#audio_file = "meeting3.wav"
		audio_file = "CIH_Test.mp3"
		#audio_file = "KIF??rDieBundeswehr.mp3"
		
		logger.info("Fixing audio file...")
		fixed_audio_file = t.fix_audio(audio_file)
		
		logger.info("Extracting full text from {}...", fixed_audio_file)
		annotated_words, full_text = t.get_words_from_text(fixed_audio_file)
		logger.info("Found {} words and text with length {}.", len(annotated_words), len(full_text))
		
		logger.info("Spell checking and correcting text...")
		repaired_text = t.repair_text(full_text)	
		logger.debug("Repaired text is {}.", repaired_text)
		
		logger.info("Calculating timestamps for corrected sentences...")
		sentences = t.get_sentences(annotated_words, repaired_text)
		
		logger.info("Getting speakers for each sentences...")
		sentences_with_speakers = t.get_speakers_from_sentences(sentences, fixed_audio_file)
		
		#logger.info("Merging sequences where possible...")
		merged_sentences = t.merge_sentences(sentences_with_speakers)
		
		logger.info("Writing document...")
		file_name, file_extension = os.path.splitext(audio_file)
		#t.write_docx(sentences_with_speakers, "{}.docx".format(file_name))
		t.write_docx(merged_sentences, "{}.docx".format(file_name))
		
		os.remove(fixed_audio_file)
		
		end = timer()
		logger.info("Done in {}", timedelta(seconds=end-start))
	