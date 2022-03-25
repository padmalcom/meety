import copy
import json
import math
import os
import sys
import textwrap
import wave
from datetime import datetime
from datetime import timedelta
from difflib import SequenceMatcher
from timeit import default_timer as timer

import numpy as np
import streamlit as st
import vosk
from docx import Document
from docx.enum.text import WD_UNDERLINE
from loguru import logger
from pydub import AudioSegment
from spacy.lang.de import German
from stqdm import stqdm
from transformers import pipeline
from webvtt import WebVTT, Caption

from recasepunc import CasePuncPredictor


# todo set tags in text block

# This can be a word or a sentence
class AnnotatedSequence:
	def __init__(self, sequence, start, end, speaker, confidence, min_confidence):
		self.sequence = sequence
		self.start = start
		self.end = end
		self.speaker = speaker
		self.confidence = confidence
		self.min_confidence = min_confidence
		
	def get_formated_start(self):
		return self.timeString(self.start)
		
	def get_formated_end(self):
		return self.timeString(self.end)
		
	def timeString(self, seconds):
		minutes = seconds / 60
		seconds = seconds % 60
		hours = int(minutes / 60)
		minutes = int(minutes % 60)
		return '%i:%02i:%06.3f' % (hours, minutes, seconds)
		
	def __repr__(self):
		return "(sequence: {}, start: {}, end: {}, speaker: {}, confidence: {}, min_confidence: {}".format(self.sequence, self.get_formated_start(), self.get_formated_end(), self.speaker, self.confidence, self.min_confidence)

class Transcriber:
	
	def __init__(self, voice_model="klein und schnell"):
		vosk.SetLogLevel(-2)
		
		if voice_model == "klein und schnell":
			model_path = 'model/vosk-model-small-de-0.15'
		else:
			model_path = 'vosk-model-de-0.21'
		spk_model_path = 'model/vosk-model-spk-0.4'
		self.COSINE_DIST = 0.4
		self.UNDETECTED_SPEAKER = "undetected speaker"
		self.UNKNOWN_SPEAKER = "unknown speaker"
		self.sample_rate = 16000
		self.model = vosk.Model(model_path)
		self.spk_model = vosk.SpkModel(spk_model_path)
		self.rec = vosk.KaldiRecognizer(self.model, self.sample_rate, self.spk_model)
		
		self.speakers={}
		self.speakers["jonas"] = [-0.849626, 0.1394, 0.377193, 0.04026, -0.516814, 0.684988, 0.276028, 0.44519, -0.071336, 0.236369, 1.438427, -1.16203, 0.171904, -0.246187, 1.162354, 1.754254, -0.870339, 0.581138, 0.277063, -0.213094, -0.178992, 1.307812, 0.383807, -1.808675, -0.183495, 0.595374, 0.552995, 1.130693, 1.509223, 0.55946, 0.377448, -0.874153, -0.805022, -0.053774, -0.024277, 0.100899, -0.16061, 0.979247, -0.806801, 1.075389, 0.054593, 0.457873, -0.123776, 0.102562, -1.03972, -0.405821, -2.296902, 1.552534, 1.987395, -0.178972, -1.699092, -1.007383, -0.044273, 0.668422, -0.387606, 1.032058, 1.218006, -0.004968, -2.110302, -0.588267, -1.630525, -0.486805, -0.149978, -0.323995, 0.057739, -1.591965, -0.507674, 2.15651, 0.087731, 0.380945, 1.659711, 1.078468, -0.96696, 1.021768, -0.429574, 1.221862, -1.227425, -0.760518, 0.883041, 1.018671, 1.070466, -0.274897, 0.163634, -0.747442, -0.05772, 0.477834, -2.581385, 0.522382, -1.506052, 1.325543, -0.888436, -0.042877, -0.999906, -0.163342, -0.045787, -2.195781, 1.082597, -0.170146, -0.511741, 0.051157, -1.417514, -1.123549, 0.37979, -0.715364, 0.868087, 0.182235, 1.849722, 2.322327, 0.353243, 0.963822, -1.506676, 1.008916, 0.883645, -0.744136, -0.141038, -0.231468, 0.424568, 0.640757, 1.176899, 0.703505, -1.75687, -0.789799, -0.169159, -1.980155, -0.769657, 0.923687, -0.850422, 1.737328]
		self.speakers["fabian"] = [-0.797457, 0.893783, -0.458962, 0.439499, 0.121801, 0.269168, -0.605457, -0.07904, 1.278862, -0.47834, 2.048105, -1.236233, 0.979613, -0.267615, 0.159861, 0.055959, -0.256166, 1.260095, 2.314549, -2.207093, -0.710628, -0.170672, 0.320985, -0.492729, 0.395429, 0.310763, -0.650441, -0.726839, 1.783956, -0.514899, -0.771253, -1.798682, -1.846042, 0.154697, 0.979613, 0.138115, 0.229337, 0.092984, -0.169231, 0.498774, -1.416122, 0.110558, 0.30648, -0.467408, 0.496537, -2.26419, 0.428986, -1.017538, 1.711265, 0.492168, -0.702733, -0.271146, -0.805249, 0.425988, 1.871118, 0.590559, 0.250679, 0.898046, 1.574314, 1.043915, -0.886984, 0.12249, 1.196055, 1.362341, 0.738961, 1.07415, 0.256881, 1.817265, -0.546605, -0.569249, -0.469055, 0.563007, -1.593265, 1.247733, 1.711871, -0.320994, -0.566991, -2.293762, 1.308869, 1.3637, -0.170194, 0.640718, -0.016229, -0.805548, -0.511637, 0.785791, -1.111833, -0.521746, -0.939944, 1.332535, 0.793953, -0.004191, -0.06619, -1.196705, 0.919897, 0.844841, 0.471952, 0.559364, 1.94756, -1.286218, -0.334231, -2.265524, 0.673755, 0.234436, -0.412774, 0.885559, 1.116514, 2.089988, -1.357796, -0.271067, 0.642276, 0.949255, 0.347166, -0.594057, 1.056821, 0.960383, -1.361658, -0.283667, -0.764461, 0.674468, -0.529306, -0.359495, 0.656774, -0.703565, 0.369017, 0.894422, -1.642501, 0.963191]
		self.new_speaker_index = 0
		
		punc_predict_path = os.path.abspath('model/vosk-recasepunc-de-0.21/checkpoint')
		self.casePuncPredictor = CasePuncPredictor(punc_predict_path, lang="de")
		
		#self.summarizer = pipeline("summarization", model="ml6team/mt5-small-german-finetune-mlsum", tokenizer="ml6team/mt5-small-german-finetune-mlsum")
		#self.summarizer = pipeline("summarization", model="T-Systems-onsite/mt5-small-sum-de-en-v2",
		# tokenizer="T-Systems-onsite/mt5-small-sum-de-en-v2")

		self.summarizer = pipeline("summarization", model="model/mt5-small-sum-de-en-v2", 
								   tokenizer="model/mt5-small-sum-de-en-v2")

		logger.info("Initialization done.")
		
	def repair_text(self, text):
		tokens = list(enumerate(self.casePuncPredictor.tokenize(text)))
		text = ''
		for token, case_label, punc_label in self.casePuncPredictor.predict(tokens, lambda x: x[1]):
			mapped = self.casePuncPredictor.map_punc_label(self.casePuncPredictor.map_case_label(token[1], case_label), punc_label)
			logger.trace("Token {}, case_label {}, punc_label {}, mapped {}", token, case_label, punc_label, mapped)
			if token[1].startswith('##'):
				text += mapped
			else:
				text +=  ' ' + mapped
		return text.strip()
		
	def get_speaker(self, fingerprint):
		best_speaker = ""
		if len(fingerprint) == 0:
			return best_speaker
		min_cd = 1000
		for s in self.speakers.keys():
			cd = self.cosine_dist(self.speakers[s], fingerprint)
			if cd < min_cd and cd < self.COSINE_DIST:
					best_speaker = s
					min_cd = cd
		if best_speaker.strip() == "":
			self.speakers[self.UNKNOWN_SPEAKER+str(self.new_speaker_index)] = fingerprint
			best_speaker = self.UNKNOWN_SPEAKER+str(self.new_speaker_index)
			self.new_speaker_index += 1
		return best_speaker, min_cd

	def cosine_dist(self, x, y):
		nx = np.array(x)
		ny = np.array(y)
		return 1 - np.dot(nx, ny) / np.linalg.norm(nx) / np.linalg.norm(ny)
		
	def get_speaker_from_audio_segment(self, audio_file, start, stop):
		logger.debug("Split audio {} from {} to {}", audio_file, start, stop)
		
		self.rec = vosk.KaldiRecognizer(self.model, self.sample_rate, self.spk_model)
		self.rec.SetWords(False)
		audio = AudioSegment.from_wav(audio_file)
		excerpt = audio[int(start*1000):int(stop*1000)]
		logger.debug("Length of excerpt is {} ms.", len(excerpt))

		timestamp = datetime.now().microsecond
		excerpt.export(str(timestamp) + ".wav", format='wav')
		wf = wave.open(str(timestamp) + ".wav", "rb")
		while True:
		#for i in tqdm(range(0, math.ceil(wf.getnframes() / 4000))):
			data = wf.readframes(4000)
			if len(data) == 0:
				break
			if self.rec.AcceptWaveform(data):
				pass
		res = self.rec.FinalResult()
		jres = json.loads(res)
		wf.close()
		os.remove(str(timestamp) + ".wav")
		if 'spk' in jres:
			speaker_name, cd = self.get_speaker(jres['spk'])
			logger.trace("Speaker fingerprint {}.", jres['spk'])
			logger.info("Speaker for {} is {}.", str(timestamp) + ".wav", speaker_name)
			return speaker_name, cd
		else:
			return self.UNDETECTED_SPEAKER, 1000
	
	def get_words_from_text(self, audio_file):
		rec_results = []
		self.rec = vosk.KaldiRecognizer(self.model, self.sample_rate, self.spk_model)
		self.rec.SetWords(True)
		wf = wave.open(audio_file, "rb")
		for i in stqdm(range(0, math.ceil(wf.getnframes() / 4000)), desc="Erkenne Audiosequenz..."):
		#while True:
		   data = wf.readframes(4000)
		   if len(data) == 0:
			   break
		   if self.rec.AcceptWaveform(data):
			   rec_results.append(self.rec.Result())
		rec_results.append(self.rec.FinalResult())
		wf.close()
		word_results = []
		full_text = ""
		for i, res in enumerate(rec_results):
			words = json.loads(res).get('result')
			if not words:
				continue
			for w in words:
				word_results.append(AnnotatedSequence(w['word'], w['start'], w['end'], None, w['conf'], w['conf']))
				logger.trace("Found new word {}.", word_results[len(word_results)-1])
				full_text += ' ' + w['word']
				
		return word_results, full_text.strip()
				
	def best_sentence_fit(self, unrepaired_words, unrepaired_start_index, repaired_words): # new
		
		punctuation = '''!()-[]{};:'"\,<>./?@#$%^&*_~'''
		
		repaired_words_without_punct = [w for w in repaired_words if not w in punctuation]

		repaired_sentence_without_punct = ' '.join(repaired_words_without_punct)

		# create i sentences
		best_ratio = 0
		best_end_index = 0
		for i in range(unrepaired_start_index, len(unrepaired_words)):
			sub_list = [x.sequence for x in unrepaired_words[unrepaired_start_index:i+1]]
			sentence_to_compare = ' '.join(sub_list)
			ratio = SequenceMatcher(None, repaired_sentence_without_punct, sentence_to_compare).ratio()
			logger.trace("{}: sentence to compare: {} ratio: {}", i, sentence_to_compare, ratio)
			if ratio > best_ratio:
				best_ratio = ratio
				best_end_index = i
				
			# Experimental - don't compare sentences, if the length of the constructed length is twice the size
			if len(repaired_sentence_without_punct)*2 < len(sentence_to_compare):
				logger.trace("Length of sentence to compare is twice as big - break.")
				break
		logger.debug("Best end {}, best ratio {}.", best_end_index, best_ratio)
		return best_end_index		
		
	def get_sentences(self, words, repaired_text):
		sentences = []
		
		# split repaired_text in sentences
		nlp = German()
		nlp.add_pipe('sentencizer')
		doc = nlp(repaired_text)
		repaired_text_sentences = [str(sent).strip() for sent in doc.sents]
		
		word_index = 0
		for rts in repaired_text_sentences:	
			words_in_sentence = [token.text for token in nlp(rts)]
			end_index = self.best_sentence_fit(words, word_index, words_in_sentence)
			confidences = [x.confidence for x in words[word_index:end_index]]
			avg_conf = 0
			min_conf = 0
			if len(confidences):
				avg_conf = sum(confidences) / (len(confidences) + 0.00000001)
				min_conf = min(confidences)
			
			#print("Lenth {}, start_index {}, end_index {}".format(len(words), word_index, end_index))
			
			sentences.append(AnnotatedSequence(rts, words[word_index].start, words[end_index].end, None, avg_conf, min_conf))
			word_index = end_index + 1
		return sentences
		
	# merge sentences with same speakers
	def merge_sentences(self, annotated_sentences):
		results = []
		current_speaker = ""
		current_sentence = ""
		current_start = 0
		current_end = 0
		confidences = []
		for i, anse in enumerate(annotated_sentences):
			if i == 0:
				current_speaker = anse.speaker
				current_sentence = anse.sequence
				current_start = anse.start
				current_end = anse.end
				confidences.append(anse.confidence)
				continue
		
			if anse.speaker != current_speaker and len(current_sentence) > 0:
				avg_conf = 0
				min_conf = 0			
				if len(confidences):
					avg_conf = sum(confidences) / (len(confidences) + 0.00000001)
					min_conf = min(confidences)
				results.append(AnnotatedSequence(current_sentence, current_start, current_end, current_speaker, avg_conf, min_conf))
				current_speaker = anse.speaker
				current_sentence = anse.sequence
				current_start = anse.start
				current_end = 0
				confidences.clear()
				confidences.append(anse.confidence)
			else:
				current_sentence += ' ' + anse.sequence
				current_end = anse.end
				confidences.append(anse.confidence)
		
		# Add last item if not empty
		if len(current_sentence) > 0:
			avg_conf = 0
			min_conf = 0		
			if len(confidences):
				avg_conf = sum(confidences) / (len(confidences) + 0.00000001)
				min_conf = min(confidences)		
			results.append(AnnotatedSequence(current_sentence, current_start, current_end, current_speaker, avg_conf, min_conf))
				
		return results
		
		
	def get_speakers_from_sentences(self, annotated_sentences, audio_file):
		results = []
		for s in stqdm(annotated_sentences, desc="Erkenne Sprecher..."):
			logger.debug("Getting speaker for {}.", s)
			speaker_name, distance = self.get_speaker_from_audio_segment(audio_file, s.start, s.end)
			logger.debug("Found speaker {}.", speaker_name)
			new_sentence = copy.copy(s)
			new_sentence.speaker = speaker_name
			results.append(new_sentence)
		return results
		
	def fix_audio(self, audio_file):
		file_name, file_extension = os.path.splitext(audio_file)
		audio = None
		is_video = False
		if file_extension.lower() == '.mp3':
			audio = AudioSegment.from_mp3(audio_file)
		elif file_extension.lower() == '.wav':
			audio = AudioSegment.from_wav(audio_file)
		elif file_extension.lower() == '.m4a':
			audio = AudioSegment.from_file(audio_file)
		elif file_extension.lower() == '.mp4':
			audio = AudioSegment.from_file(audio_file)
			is_video = True
		else:
			logger.error("File format {} not supported.", file_extension)
			sys.exit(0)
			
		audio = audio.set_channels(1)
		audio = audio.set_frame_rate(16000)
		audio = audio.set_sample_width(2)
		timestamp = datetime.now().microsecond
		file_name = "conv" + str(timestamp) + ".wav"
		audio.export(file_name, format='wav', bitrate="64k")
		return file_name, is_video
		
	def write_docx(self, title, annotated_sentences, out_file):
		document = Document()
		document.add_heading(title, 0)
		
		p = None
		last_speaker = None
		for anse in annotated_sentences:
			if anse.speaker != last_speaker or p == None:
				p = document.add_paragraph('')
				if show_timestamp:
					p.add_run("[{} - {}] {}:".format(anse.start, anse.end, anse.speaker)).bold = True
				else:
					p.add_run("{}:".format(anse.speaker)).bold = True

			if underline_uncertainties and anse.min_confidence < 0.5:
				p.add_run(' ' + anse.sequence).underline = WD_UNDERLINE.WAVY
			else:
				p.add_run(' ' + anse.sequence).underline = False
								
			last_speaker = anse.speaker

		document.save(out_file)
		
	def write_webvtt(self, annotated_sentences, out_file):
		vtt = WebVTT()
		for anse in annotated_sentences:
			caption = Caption(anse.get_formated_start(), anse.get_formated_end(), textwrap.fill(anse.speaker + ": " + anse.sequence))
			vtt.captions.append(caption)
		vtt.save(out_file)
		
		
	def summarize(self, text):
		min_length = max(5, int(len(text) * 0.05))
		max_length = min(200, int(len(text) * 0.3))
		s = self.summarizer(text, min_length=min_length, max_length=max_length)
		return s[0]['summary_text']
		
def create_transcript(temp_file_name, original_file_name, original_file_extension, show_log, summary, underline_uncertainties, show_timestamp, ai_strong):
	start = timer()
	t = Transcriber(voice_model=ai_strong)
	audio_file = temp_file_name
	text = ""
	placeholder = st.empty()
	
	if show_log:
		end = timer()
		text += "- __[{}]__ Preparing file...  \n".format(timedelta(seconds=end-start))
		placeholder.info(text)
	fixed_audio_file, is_video = t.fix_audio(audio_file)
	
	if show_log:
		end = timer()
		if is_video:
			text += "- __[{}]__ Oh look, it's a video! :tv:  \n".format(timedelta(seconds=end-start))
		else:
			text += "- __[{}]__ Oh look, it's an audio file! :loudspeaker:  \n".format(timedelta(seconds=end-start))
		placeholder.info(text)
	
	if show_log:
		end = timer()
		text += "- __[{}]__ Extracting full text from prepared file {}...  \n".format(timedelta(seconds=end-start), fixed_audio_file)
		placeholder.info(text)
	annotated_words, full_text = t.get_words_from_text(fixed_audio_file)
	for aw in annotated_words:
		print(aw)
	
	if show_log:
		end = timer()
		text += "- __[{}]__ Found **{} words** and text with **length {}** :newspaper:.  \n".format(timedelta(seconds=end-start), len(annotated_words), len(full_text))
		placeholder.info(text)
	
	if show_log:
		end = timer()
		text += "- __[{}]__ Spell checking and correcting text...  \n".format(timedelta(seconds=end-start))
		placeholder.info(text)
	repaired_text = t.repair_text(full_text)	
	
	if show_log:
		end = timer()
		text += "- __[{}]__ Text has been corrected by a highly over-qualified monkey :monkey_face:  \n".format(timedelta(seconds=end-start), repaired_text)
		placeholder.info(text)
	
	if show_log:
		end = timer()
		text += "- __[{}]__ Calculating timestamps for corrected sentences... :watch:  \n".format(timedelta(seconds=end-start))
		placeholder.info(text)
	sentences = t.get_sentences(annotated_words, repaired_text)
	
	if show_log:
		end = timer()
		text += "- __[{}]__ Getting speakers for each sentences... :speech_balloon:  \n".format(timedelta(seconds=end-start))
		placeholder.info(text)
	sentences_with_speakers = t.get_speakers_from_sentences(sentences, fixed_audio_file)
	
	summarization = ""
	if summary:
		end = timer()
		text += "- __[{}]__ Summarizing...  \n".format(timedelta(seconds=end-start))
		placeholder.info(text)
		
		full_text = ' '.join([swp.sequence for swp in sentences_with_speakers])
		summarization = t.summarize(full_text)
		st.markdown("**tl;dr:** {}".format(summarization))

	
	last_speaker = None
	output_text = ""
	for anse in sentences_with_speakers:
		if anse.speaker != last_speaker:
			if len(output_text) > 0:
				st.markdown(output_text)
				output_text = ""
			if show_timestamp:
				output_text += "[{} - {}] **{}**: ".format(anse.start, anse.end, anse.speaker)
			else:
				output_text += "**{}**: ".format(anse.speaker)
				
		if underline_uncertainties and anse.min_confidence < 0.6:
			output_text += ' _' + anse.sequence + '_'
		else:
			output_text += ' _' + anse.sequence + '_'
		last_speaker = anse.speaker
	if len(output_text) > 0:
		st.markdown(output_text)
			
	if show_log:
		end = timer()
		text += "- __[{}]__ Writing output... :pencil2:  \n".format(timedelta(seconds=end-start))
		placeholder.info(text)
	file_name, file_extension = os.path.splitext(audio_file)
	
	os.remove(fixed_audio_file)
	
	end = timer()
	
	st.success("Done in {}  \n".format(timedelta(seconds=end-start)))
	
	# generate docx
	docx_filename = original_file_name + ".docx"
	t.write_docx(original_file_name, sentences_with_speakers, docx_filename)
	with open(docx_filename, 'rb') as f:
		st.download_button('Download Trankript (*.docx)', f, file_name=docx_filename)
		
	# generate subtitles
	if is_video:
		webvtt_filename = original_file_name + ".vtt"
		t.write_webvtt(sentences_with_speakers, webvtt_filename)
		with open(webvtt_filename, 'rb') as f1:
			st.download_button('Download subtitles (*.vtt)', f1, file_name=webvtt_filename)
	st.balloons()
	
			
if __name__ == '__main__':
	logger.remove()
	logger.add(sys.stderr, level="INFO")
	
	st.title("Meety - Herzlichen Willkommen!")
	st.header("AI-basierte Audiotranskription für Web-Konferenzen")
	
	st.markdown("Diese Anwendung hilft dir, Meetings automatisiert zu verschriftlichen und dabei zwischen verschiedenen Sprechnern zu unterscheiden. Neugierig geworden? Probier es mal aus!")
	
	summary = st.checkbox("Generiere eine Zusammenfassung des Meetings.", value=True)
	underline_uncertainties = st.checkbox("Schreibe Sätze mit hoher Ungewissheit kursiv.", value=True)
	show_timestamp = st.checkbox("Zeige Zeitstempel der Sätze.", value=True)
	ai_strong = st.radio("Starke KI nutzen (langsamer, aber genauer)?", ('groß und langsam', 'klein und schnell'), index=1)
	
	show_log = st.checkbox("Zeig mir was genau du machst!", value=True)
		
	temp_file_name = ""
	original_file_name = ""
	original_file_extension = ""
	
	st.markdown("Grad' kein Meetingmitschnitt zuhand? Lade einfach unsere kurze Demo!")
	with open("demo.mp3", 'rb') as f:
		st.download_button('Download demo.mp3', f, file_name="demo.mp3")
		
	uploaded_file = st.file_uploader("Wählen eine Audiodatei", type=['wav', 'mp3', 'm4a', 'mp4'])
	if uploaded_file is not None:

		# store file as temp file
		timestamp = datetime.now().microsecond		
		original_file_name, original_file_extension = os.path.splitext(uploaded_file.name)
		temp_file_name = str(timestamp) + original_file_extension
		
		# write uploaded file
		with open(temp_file_name, "wb") as f:
			f.write(uploaded_file.getbuffer())
			st.success("Datei {} erfolgreich hochgeladen.".format(uploaded_file.name))		

			
			st.button("Erstelle Transkript", on_click=create_transcript, args=(temp_file_name, original_file_name, original_file_extension, show_log, summary, underline_uncertainties, show_timestamp, ai_strong, ))
		